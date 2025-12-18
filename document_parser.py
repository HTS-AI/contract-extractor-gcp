"""
Document Parser for various contract document formats.
Supports PDF, DOCX, OCR, and plain text inputs.
"""

import os
import json
from typing import Optional, List, Dict, Any
from pathlib import Path

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import pymupdf  # PyMuPDF for better table extraction
    try:
        import pandas as pd
        PANDAS_AVAILABLE = True
    except ImportError:
        PANDAS_AVAILABLE = False
        pd = None
except ImportError:
    pymupdf = None
    PANDAS_AVAILABLE = False
    pd = None

try:
    from docx import Document
except ImportError:
    Document = None

# Local OCR dependencies removed - using Google Cloud Vision API only

try:
    from vision_gcp import vision_ocr_pdf
    from gcs_utils import upload_file_to_gcs, read_text_from_gcs
    VISION_API_AVAILABLE = True
except (ImportError, ModuleNotFoundError) as e:
    # Log the actual error for debugging
    import sys
    print(f"Warning: Could not import Vision API modules: {e}", file=sys.stderr)
    VISION_API_AVAILABLE = False
    vision_ocr_pdf = None
    upload_file_to_gcs = None
    read_text_from_gcs = None
except Exception as e:
    # Catch any other import-related errors
    import sys
    print(f"Warning: Unexpected error importing Vision API modules: {e}", file=sys.stderr)
    VISION_API_AVAILABLE = False
    vision_ocr_pdf = None
    upload_file_to_gcs = None
    read_text_from_gcs = None

# Image enhancement removed - only using OCR and regular PDF parsing


class DocumentParser:
    """Parses contract documents from various formats."""
    
    def __init__(
        self, 
        use_gcs_vision: bool = True,
        gcs_input_path: str = "gs://data-pdf-extractor/input-docs/",
        gcs_output_path: str = "gs://data-pdf-extractor/processed-documents/",
        gcs_extracted_text_path: str = "gs://data-pdf-extractor/extracted-text/",
        service_account_file: Optional[str] = None
    ):
        """
        Initialize the document parser.
        
        Args:
            use_gcs_vision: If True, use Google Cloud Vision API for scanned PDFs
            gcs_input_path: GCS path for input files
            gcs_output_path: GCS path for Vision API JSON outputs
            gcs_extracted_text_path: GCS path for extracted text files
            service_account_file: DEPRECATED - Credentials are loaded from GCP_CREDENTIALS_JSON environment variable
        """
        # Re-check Vision API availability at runtime (in case modules were loaded after document_parser)
        runtime_vision_available = VISION_API_AVAILABLE
        if not runtime_vision_available:
            # Try importing again at runtime
            try:
                from vision_gcp import vision_ocr_pdf
                from gcs_utils import upload_file_to_gcs, read_text_from_gcs
                runtime_vision_available = True
                # Update module-level variables
                import sys
                if 'vision_gcp' in sys.modules and 'gcs_utils' in sys.modules:
                    globals()['vision_ocr_pdf'] = vision_ocr_pdf
                    globals()['upload_file_to_gcs'] = upload_file_to_gcs
                    globals()['read_text_from_gcs'] = read_text_from_gcs
            except Exception as e:
                runtime_vision_available = False
        
        self.use_gcs_vision = use_gcs_vision and runtime_vision_available
        if use_gcs_vision and not runtime_vision_available:
            import warnings
            warnings.warn(
                "use_gcs_vision=True was requested but Vision API is not available. "
                "Please ensure vision_gcp.py and gcs_utils.py are in the Python path.",
                UserWarning
            )
        
        self.gcs_input_path = gcs_input_path.rstrip('/') + '/'
        self.gcs_output_path = gcs_output_path.rstrip('/') + '/'
        self.gcs_extracted_text_path = gcs_extracted_text_path.rstrip('/') + '/'
        
        # Credentials are now loaded from GCP_CREDENTIALS_JSON environment variable
        # Verify the environment variable is set
        if service_account_file:
            import warnings
            warnings.warn(
                "service_account_file parameter is deprecated. "
                "Using GCP_CREDENTIALS_JSON from environment instead.",
                DeprecationWarning
            )
        
        # Verify GCP_CREDENTIALS_JSON is available (but don't parse/print its value)
        credentials_json = os.getenv('GCP_CREDENTIALS_JSON')
        if not credentials_json:
            print("[WARNING] GCP_CREDENTIALS_JSON environment variable is not set")
            print("[WARNING] OCR will fail if credentials are not available via environment variables")
        elif not credentials_json.strip():
            print("[WARNING] GCP_CREDENTIALS_JSON environment variable is empty")
            print("[WARNING] OCR will fail if credentials are not available via environment variables")
        else:
            print("[INFO] GCP_CREDENTIALS_JSON is available (credentials loaded from environment)")
    
    def parse(self, file_path: str, use_ocr: bool = False) -> str:
        """
        Parse a document and extract text.
        
        Args:
            file_path: Path to the document file
            use_ocr: If True, use OCR for PDF files (for scanned documents)
            
        Returns:
            Extracted text from the document
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file does not exist
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            # Only use OCR/Vision API if explicitly enabled
            if self.use_gcs_vision and (use_ocr or self._is_scanned_pdf(file_path)):
                # Use GCS Vision API (required for OCR)
                try:
                    # Try importing to check availability
                    from vision_gcp import vision_ocr_pdf
                    from gcs_utils import upload_file_to_gcs, read_text_from_gcs
                    # If import succeeds, use Vision API
                    return self._parse_pdf_with_vision_api(file_path)
                except ImportError as e:
                    # If import fails, fall back to regular PDF parsing
                    import warnings
                    warnings.warn(
                        f"Vision API not available ({str(e)}). Falling back to PyPDF2 extraction.",
                        UserWarning
                    )
                    return self._parse_pdf(file_path)
            else:
                # Use regular PDF parsing (PyPDF2)
                return self._parse_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self._parse_docx(file_path)
        elif file_ext in ['.txt', '.text']:
            return self._parse_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def parse_with_pages(self, file_path: str, use_ocr: bool = False) -> tuple[str, dict]:
        """
        Parse a document and extract text with page information.
        
        Args:
            file_path: Path to the document file
            use_ocr: If True, use OCR for PDF files (for scanned documents)
            
        Returns:
            Tuple of (full_text, page_map) where:
            - full_text: Complete extracted text
            - page_map: Dictionary mapping page numbers to text content
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file does not exist
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            # Only use OCR/Vision API if explicitly enabled
            if self.use_gcs_vision and (use_ocr or self._is_scanned_pdf(file_path)):
                # Use GCS Vision API (required for OCR)
                try:
                    # Try importing to check availability
                    from vision_gcp import vision_ocr_pdf
                    from gcs_utils import upload_file_to_gcs, read_text_from_gcs
                    # If import succeeds, use Vision API
                    return self._parse_pdf_with_vision_api_with_pages(file_path)
                except ImportError as e:
                    # If import fails, fall back to regular PDF parsing
                    import warnings
                    warnings.warn(
                        f"Vision API not available ({str(e)}). Falling back to PyPDF2 extraction.",
                        UserWarning
                    )
                    return self._parse_pdf_with_pages(file_path)
            else:
                # Use regular PDF parsing (PyPDF2)
                return self._parse_pdf_with_pages(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self._parse_docx_with_pages(file_path)
        elif file_ext in ['.txt', '.text']:
            return self._parse_text_with_pages(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _parse_pdf(self, file_path: str) -> str:
        """Extract text from PDF using PyPDF2."""
        if PyPDF2 is None:
            raise ImportError("PyPDF2 is required for PDF parsing. Install it with: pip install PyPDF2")
        
        text_parts = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
        except Exception as e:
            raise ValueError(f"Error parsing PDF: {str(e)}")
        
        return '\n'.join(text_parts)
    
    def _parse_pdf_with_pages(self, file_path: str) -> tuple[str, dict]:
        """Extract text from PDF with page information."""
        if PyPDF2 is None:
            raise ImportError("PyPDF2 is required for PDF parsing. Install it with: pip install PyPDF2")
        
        text_parts = []
        page_map = {}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        page_number = page_num + 1  # 1-indexed
                        text_parts.append(text)
                        page_map[page_number] = text
        except Exception as e:
            raise ValueError(f"Error parsing PDF: {str(e)}")
        
        full_text = '\n'.join(text_parts)
        return full_text, page_map
    
    
    def _parse_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        if Document is None:
            raise ImportError("python-docx is required for DOCX parsing. Install it with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
        except Exception as e:
            raise ValueError(f"Error parsing DOCX: {str(e)}")
        
        return '\n'.join(text_parts)
    
    def _parse_docx_with_pages(self, file_path: str) -> tuple[str, dict]:
        """Extract text from DOCX file with page information (estimated)."""
        if Document is None:
            raise ImportError("python-docx is required for DOCX parsing. Install it with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            text_parts = []
            page_map = {}
            current_page = 1
            current_page_text = []
            chars_per_page = 2000  # Estimate: ~2000 characters per page
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
                    current_page_text.append(paragraph.text)
                    
                    # Estimate page breaks (rough approximation)
                    if len('\n'.join(current_page_text)) > chars_per_page:
                        page_map[current_page] = '\n'.join(current_page_text)
                        current_page += 1
                        current_page_text = [paragraph.text]
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text = ' | '.join(row_text)
                        text_parts.append(text)
                        current_page_text.append(text)
                        
                        if len('\n'.join(current_page_text)) > chars_per_page:
                            page_map[current_page] = '\n'.join(current_page_text)
                            current_page += 1
                            current_page_text = []
            
            # Add remaining text to last page
            if current_page_text:
                page_map[current_page] = '\n'.join(current_page_text)
        except Exception as e:
            raise ValueError(f"Error parsing DOCX: {str(e)}")
        
        full_text = '\n'.join(text_parts)
        return full_text, page_map
    
    def _parse_text(self, file_path: str) -> str:
        """Read plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
    
    def _parse_text_with_pages(self, file_path: str) -> tuple[str, dict]:
        """Read plain text file with page information (estimated)."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                full_text = file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                full_text = file.read()
        
        # Estimate pages (rough approximation: ~2000 chars per page)
        page_map = {}
        chars_per_page = 2000
        lines = full_text.split('\n')
        current_page = 1
        current_page_text = []
        current_chars = 0
        
        for line in lines:
            current_page_text.append(line)
            current_chars += len(line)
            
            if current_chars > chars_per_page:
                page_map[current_page] = '\n'.join(current_page_text)
                current_page += 1
                current_page_text = []
                current_chars = 0
        
        # Add remaining text to last page
        if current_page_text:
            page_map[current_page] = '\n'.join(current_page_text)
        
        return full_text, page_map
    
    def _is_scanned_pdf(self, file_path: str) -> bool:
        """
        Detect if a PDF is scanned (image-based) by checking text extraction.
        Improved detection with multiple checks.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            bool: True if PDF appears to be scanned (little/no text extracted)
        """
        if PyPDF2 is None:
            # If PyPDF2 is not available, try using pymupdf as fallback
            if pymupdf is not None:
                try:
                    doc = pymupdf.open(file_path)
                    total_text_length = 0
                    pages_to_check = min(3, len(doc))
                    
                    for page_num in range(pages_to_check):
                        page = doc[page_num]
                        text = page.get_text()
                        if text:
                            total_text_length += len(text.strip())
                    
                    doc.close()
                    avg_text_per_page = total_text_length / pages_to_check if pages_to_check > 0 else 0
                    is_scanned = avg_text_per_page < 100
                    
                    if is_scanned:
                        print(f"[SCAN_DETECTION] Detected scanned PDF (avg {avg_text_per_page:.0f} chars/page)")
                    return is_scanned
                except Exception as e:
                    print(f"[SCAN_DETECTION] Error checking with pymupdf: {e}")
                    return False
            return False
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_text_length = 0
                pages_with_text = 0
                
                # Check first few pages for text (up to 5 pages for better detection)
                pages_to_check = min(5, len(pdf_reader.pages))
                if pages_to_check == 0:
                    return False
                
                for page_num in range(pages_to_check):
                    try:
                        page = pdf_reader.pages[page_num]
                        text = page.extract_text()
                        if text and text.strip():
                            text_length = len(text.strip())
                            total_text_length += text_length
                            pages_with_text += 1
                    except Exception as e:
                        # If we can't extract from a page, it might be scanned
                        print(f"[SCAN_DETECTION] Warning: Could not extract text from page {page_num + 1}: {e}")
                
                # Multiple criteria for scanned PDF detection:
                # 1. Average text per page is very low (< 100 chars)
                # 2. Less than 50% of pages have extractable text
                avg_text_per_page = total_text_length / pages_to_check if pages_to_check > 0 else 0
                text_pages_ratio = pages_with_text / pages_to_check if pages_to_check > 0 else 0
                
                is_scanned = avg_text_per_page < 100 or text_pages_ratio < 0.5
                
                if is_scanned:
                    print(f"[SCAN_DETECTION] Detected scanned PDF:")
                    print(f"   - Average text per page: {avg_text_per_page:.0f} characters")
                    print(f"   - Pages with text: {pages_with_text}/{pages_to_check}")
                    print(f"   - Text pages ratio: {text_pages_ratio:.1%}")
                else:
                    print(f"[SCAN_DETECTION] PDF appears to have extractable text (avg {avg_text_per_page:.0f} chars/page)")
                
                return is_scanned
        except Exception as e:
            print(f"[SCAN_DETECTION] Error detecting scanned PDF: {e}")
            # If we can't determine, assume it's not scanned (safer default)
            return False
    
    def _parse_pdf_with_vision_api(self, file_path: str) -> str:
        """
        Parse scanned PDF using Google Cloud Vision API.
        
        Args:
            file_path: Path to local PDF file
            
        Returns:
            str: Extracted text from the PDF
        """
        # Try importing at runtime (modules might be available now even if not at import time)
        try:
            from vision_gcp import vision_ocr_pdf
            from gcs_utils import upload_file_to_gcs, read_text_from_gcs
        except ImportError as e:
            raise ImportError(
                f"Vision API modules not available. Error: {str(e)}\n"
                "Please ensure vision_gcp.py and gcs_utils.py are in your Python path "
                "and all required dependencies (google-cloud-vision, google-cloud-storage) are installed."
            )
        
        from pathlib import Path
        
        # Get filename
        filename = Path(file_path).name
        
        # Upload file to GCS input folder
        gcs_input_uri = f"{self.gcs_input_path}{filename}"
        print(f"Uploading {filename} to {gcs_input_uri}...")
        upload_file_to_gcs(file_path, gcs_input_uri, None)  # Uses GCP_CREDENTIALS_JSON from environment
        
        # Generate unique output folder for this document
        doc_id = Path(file_path).stem
        gcs_output_uri = f"{self.gcs_output_path}{doc_id}/"
        
        # Process with Vision API
        print(f"Processing {filename} with Google Cloud Vision API...")
        try:
            extracted_text = vision_ocr_pdf(
                gcs_input_uri,
                gcs_output_uri,
                gcs_input_path=self.gcs_input_path,
                service_account_file=None  # Uses GCP_CREDENTIALS_JSON from environment
            )
            
            if not extracted_text or len(extracted_text.strip()) == 0:
                print(f"[OCR_WARNING] Vision API completed but no text was extracted")
                print(f"[OCR_WARNING] This might indicate:")
                print(f"   - The document is completely blank")
                print(f"   - OCR processing failed silently")
                print(f"   - The document format is not supported")
                # Fall back to regular PDF extraction
                print(f"[OCR_FALLBACK] Attempting fallback to regular PDF extraction...")
                return self._parse_pdf(file_path)
            else:
                print(f"[OCR_SUCCESS] Extracted {len(extracted_text)} characters using Vision API OCR")
                return extracted_text
                
        except Exception as e:
            error_msg = str(e)
            print(f"[OCR_ERROR] Vision API OCR failed: {error_msg}")
            
            # Check for specific error types
            if "Invalid JWT Signature" in error_msg or "invalid_grant" in error_msg:
                print(f"[OCR_ERROR] Authentication failed - GCP credentials are invalid")
                print(f"[OCR_ERROR] Please check your GCP_CREDENTIALS_JSON environment variable")
            elif "PermissionDenied" in error_msg or "permission" in error_msg.lower():
                print(f"[OCR_ERROR] Permission denied - service account lacks required permissions")
            elif "timeout" in error_msg.lower():
                print(f"[OCR_ERROR] OCR processing timed out - document may be too large")
            else:
                print(f"[OCR_ERROR] Unexpected error during OCR processing")
            
            print(f"[OCR_FALLBACK] Falling back to regular PDF extraction...")
            # Fall back to regular PDF extraction
            return self._parse_pdf(file_path)
    
    def _parse_pdf_with_vision_api_with_pages(self, file_path: str) -> tuple[str, dict]:
        """
        Parse scanned PDF using Google Cloud Vision API with page information.
        
        Args:
            file_path: Path to local PDF file
            
        Returns:
            Tuple of (full_text, page_map)
        """
        # Get full text (with error handling)
        try:
            full_text = self._parse_pdf_with_vision_api(file_path)
        except Exception as e:
            # If OCR fails, fall back to regular PDF parsing
            print(f"[OCR_FALLBACK] Vision API failed, using regular PDF extraction: {e}")
            return self._parse_pdf_with_pages(file_path)
        
        # Estimate page breaks (Vision API doesn't provide exact page mapping in text)
        # Split text by double newlines or estimate based on length
        page_map = {}
        lines = full_text.split('\n\n')
        chars_per_page = 2000  # Estimate
        current_page = 1
        current_page_text = []
        current_chars = 0
        
        for line in lines:
            current_page_text.append(line)
            current_chars += len(line)
            
            if current_chars > chars_per_page:
                page_map[current_page] = '\n\n'.join(current_page_text)
                current_page += 1
                current_page_text = []
                current_chars = 0
        
        # Add remaining text to last page
        if current_page_text:
            page_map[current_page] = '\n\n'.join(current_page_text)
        
        return full_text, page_map
    
    def extract_tables(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract tables from document (PDF or DOCX).
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of tables, each with page number, headers, and rows
        """
        file_ext = Path(file_path).suffix.lower()
        tables = []
        
        if file_ext == '.pdf':
            tables = self._extract_tables_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            tables = self._extract_tables_from_docx(file_path)
        # TXT files don't have structured tables
        
        return tables
    
    def _extract_tables_from_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract tables from PDF using PyMuPDF if available."""
        tables = []
        
        # Try PyMuPDF first (better table extraction)
        if pymupdf is not None:
            try:
                doc = pymupdf.open(file_path)
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    try:
                        # Use PyMuPDF's table extraction
                        page_tables = pymupdf.table.find_tables(page)
                        for table in page_tables:
                            # Convert table to structured format
                            if PANDAS_AVAILABLE and pd is not None:
                                try:
                                    table_data = table.to_pandas()
                                    headers = table_data.columns.tolist() if not table_data.empty else []
                                    rows = table_data.values.tolist()
                                except:
                                    # Fallback to markdown if pandas fails
                                    table_md = table.to_markdown()
                                    # Parse markdown table
                                    lines = table_md.split('\n')
                                    headers = []
                                    rows = []
                                    for i, line in enumerate(lines):
                                        if '|' in line:
                                            cells = [c.strip() for c in line.split('|') if c.strip() and not all(c in '-: ' for c in c.strip())]
                                            if i == 0:
                                                headers = cells
                                            elif cells:
                                                rows.append(cells)
                            else:
                                # Fallback to markdown if pandas not available
                                table_md = table.to_markdown()
                                lines = table_md.split('\n')
                                headers = []
                                rows = []
                                for i, line in enumerate(lines):
                                    if '|' in line:
                                        cells = [c.strip() for c in line.split('|') if c.strip() and not all(c in '-: ' for c in c.strip())]
                                        if i == 0:
                                            headers = cells
                                        elif cells:
                                            rows.append(cells)
                            
                            tables.append({
                                "page": page_num + 1,
                                "headers": headers,
                                "rows": rows,
                                "row_count": len(rows),
                                "column_count": len(headers) if headers else 0
                            })
                    except Exception as e:
                        # If table extraction fails for this page, continue
                        print(f"[TABLE EXTRACTION] Warning: Could not extract tables from page {page_num + 1}: {e}")
                        continue
                doc.close()
            except Exception as e:
                print(f"[TABLE EXTRACTION] Warning: PyMuPDF table extraction failed: {e}")
        
        # Fallback: Try to detect tables from text using heuristics
        if not tables:
            try:
                document_text, page_map = self.parse_with_pages(file_path)
                # Simple heuristic: look for lines with multiple separators (|, tabs, multiple spaces)
                for page_num, page_text in page_map.items():
                    lines = page_text.split('\n')
                    potential_table_rows = []
                    
                    for line in lines:
                        # Check if line looks like a table row (has multiple separators)
                        if '|' in line and line.count('|') >= 2:
                            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                            if len(cells) >= 2:
                                potential_table_rows.append(cells)
                        elif '\t' in line and line.count('\t') >= 1:
                            cells = [cell.strip() for cell in line.split('\t') if cell.strip()]
                            if len(cells) >= 2:
                                potential_table_rows.append(cells)
                    
                    # If we found multiple rows with similar structure, it's likely a table
                    if len(potential_table_rows) >= 2:
                        # Use first row as headers if it looks like headers
                        headers = potential_table_rows[0] if potential_table_rows else []
                        rows = potential_table_rows[1:] if len(potential_table_rows) > 1 else []
                        
                        tables.append({
                            "page": page_num,
                            "headers": headers,
                            "rows": rows,
                            "row_count": len(rows),
                            "column_count": len(headers) if headers else 0,
                            "extraction_method": "heuristic"
                        })
            except Exception as e:
                print(f"[TABLE EXTRACTION] Warning: Heuristic table extraction failed: {e}")
        
        return tables
    
    def _extract_tables_from_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract tables from DOCX file."""
        tables = []
        
        if Document is None:
            return tables
        
        try:
            doc = Document(file_path)
            current_page = 1
            chars_per_page = 2000  # Estimate pages
            
            for table in doc.tables:
                headers = []
                rows = []
                
                # Extract table data
                for i, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                    
                    if i == 0:
                        # First row as headers
                        headers = row_data
                    else:
                        rows.append(row_data)
                
                if headers or rows:
                    tables.append({
                        "page": current_page,  # Estimated page
                        "headers": headers,
                        "rows": rows,
                        "row_count": len(rows),
                        "column_count": len(headers) if headers else 0,
                        "extraction_method": "docx_native"
                    })
                
                # Estimate page increment (rough approximation)
                total_chars = sum(len(' '.join(row)) for row in [headers] + rows)
                if total_chars > chars_per_page:
                    current_page += 1
                    
        except Exception as e:
            print(f"[TABLE EXTRACTION] Error extracting tables from DOCX: {e}")
        
        return tables

